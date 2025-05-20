from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Word document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Add Name and Contact Info
doc.add_heading('Ashutosh Pradhan', level=1)
contact_info = doc.add_paragraph()
contact_info.add_run('Location • Phone Number • Email Address\n')
contact_info.add_run('LinkedIn: ').bold = True
contact_info.add_run('www.linkedin.com/in/ashutosh-pradhan-developer\n')
contact_info.add_run('GitHub: ').bold = True
contact_info.add_run('https://github.com/Ashutosh-kapsus | https://github.com/ashutosh164')

# Section: Professional Summary
doc.add_heading('Professional Summary', level=2)
doc.add_paragraph(
    "Results-driven Python Full Stack Developer with 3.5+ years of experience in designing, developing, and deploying "
    "scalable web applications using Django, Django Rest Framework, and JavaScript frameworks like Nuxt.js and React. "
    "Passionate about building high-performance applications and exploring data science and AI/ML technologies."
)

# Section: Technical Skills
doc.add_heading('Technical Skills', level=2)
doc.add_paragraph(
    "- Languages: Python, JavaScript, SQL, HTML, CSS, Shell Script\n"
    "- Frameworks/Libraries: Django, DRF, React, Nuxt3, Bootstrap, Pandas, NumPy, PySpark\n"
    "- Databases: PostgreSQL, MySQL, MongoDB\n"
    "- Tools/Technologies: Git, Docker, Jenkins, AWS, Linux\n"
    "- Big Data & Cloud: PySpark, Hive, HDFS, AWS EC2, S3\n"
    "- Others: REST APIs, MVC/MVT, Agile, Jira"
)

# Section: Professional Experience
doc.add_heading('Professional Experience', level=2)

doc.add_paragraph('Python Full Stack Developer\nCompany Name – Location | Month Year – Present', style='List Bullet')
doc.add_paragraph(
    "• Developed RESTful APIs and web applications using Django and DRF, reducing API latency by 30%.\n"
    "• Built dynamic front-end interfaces with Nuxt.js and integrated seamlessly with back-end APIs.\n"
    "• Collaborated with cross-functional teams to deploy applications on AWS EC2 and manage data on S3.\n"
    "• Implemented role-based access control and JWT authentication for secure user management.\n"
    "• Created unit tests using PyTest, increasing code coverage to 85%."
)

doc.add_paragraph('Software Developer Intern\nCompany Name – Location | Month Year – Month Year', style='List Bullet')
doc.add_paragraph(
    "• Assisted in developing data ingestion pipelines using PySpark and Hive on Hadoop.\n"
    "• Wrote efficient SQL queries for data extraction and transformation from large datasets.\n"
    "• Contributed to system deployment scripts using Shell scripting and Docker."
)

# Section: Projects
doc.add_heading('Projects', level=2)

doc.add_paragraph('AI-Based Recommendation System', style='List Bullet')
doc.add_paragraph(
    "• Built a recommendation engine using scikit-learn and Pandas to improve user personalization.\n"
    "• Integrated with Django backend and deployed on Heroku."
)

doc.add_paragraph('E-commerce Web Application – https://new-eu.shoptsy.in/shop', style='List Bullet')
doc.add_paragraph(
    "• Developed a full-stack web app using Django and React, including cart, order, and payment modules.\n"
    "• Implemented PostgreSQL as the database and used AWS S3 for media storage."
)

doc.add_paragraph('Forex Cashback Platform – https://www.forexcashbackrebate.com/', style='List Bullet')
doc.add_paragraph(
    "• Developed user dashboards, cashback tracking, and automated email notifications.\n"
    "• Built backend modules with Django and integrated real-time transaction data feeds."
)

# Section: Education
doc.add_heading('Education', level=2)
doc.add_paragraph('Bachelor of Technology in Computer Science\nUniversity Name – Location | Month Year – Month Year')

# Section: Certifications
doc.add_heading('Certifications', level=2)
doc.add_paragraph('- [Certification Name], [Issuing Organization], [Year]\n- [Certification Name], [Issuing Organization], [Year]')

# Save the document as a PDF-ready Word file (will convert to PDF later)
doc_path = "/mnt/data/Ashutosh_Pradhan_ATS_Resume.docx"
doc.save(doc_path)

doc_path
